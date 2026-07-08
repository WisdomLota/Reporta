"""
YOOWA Weekly Report Engine
--------------------------
Reads the 'Lagos 2025' workbook and fills the Weekly Accounting Report .docx.

Confirmed mappings (validated against the 08-14 June report):
  Lagos (Lefkosa)  -> sheet 'HAMITKOY SALES'
  Abuja (Magusa)   -> sheet 'MAGUSA SALES'
  Sales tabs: column B = date. Money columns located by header text on row 2:
      NAIRA, CASH, POS + IBAN, USDT, REDOT, USD
  EXPENSE SHEET: row 1 = dates; each date spans 2 sub-columns
      (row 3 sub-header) HAMITKOY | MAGUSA. Items in column A, category col B.

Anything uncertain (exchange rate, USDT/USD/foreign rates) is passed in as
parameters from the UI so the user confirms once, not guessed by the code.
"""

import datetime
import io
import re
from copy import deepcopy

import openpyxl
from docx import Document


# ----------------------------------------------------------------------------
# Helpers
# ----------------------------------------------------------------------------
def _as_date(v):
    if isinstance(v, datetime.datetime):
        return v.date()
    if isinstance(v, datetime.date):
        return v
    return None


def _num(v):
    """Coerce a cell into a float, treating blanks/None/text as 0."""
    if v is None:
        return 0.0
    if isinstance(v, (int, float)):
        return float(v)
    s = str(v).strip().replace(",", "")
    if s in ("", "-"):
        return 0.0
    try:
        return float(s)
    except ValueError:
        return 0.0


def _money(v):
    """Format a number the way the report does: ₺ 1,234.56 (trim trailing .00)."""
    v = round(float(v), 2)
    if v == int(v):
        return f"₺ {int(v):,}"
    return f"₺ {v:,.2f}"


def _find_money_columns(ws):
    """Locate money columns by scanning row 2 (and row 3 fallback) header text."""
    wanted = {
        "NAIRA": ["NAIRA"],
        "CASH": ["CASH"],
        "POS": ["POS + IBAN", "POS+IBAN", "POS"],
        "USDT": ["USDT"],
        "REDOT": ["REDOT"],
        "USD": ["USD"],
    }
    found = {}
    for col in range(1, ws.max_column + 1):
        for hdr_row in (2, 3):
            raw = ws.cell(hdr_row, col).value
            if not raw:
                continue
            txt = str(raw).strip().upper()
            for key, variants in wanted.items():
                if key in found:
                    continue
                if any(txt == v.upper() for v in variants):
                    found[key] = col
    return found


def _table_by_header(doc, header_text):
    """Find the first table whose top-left or any first-column cell contains header_text."""
    for t in doc.tables:
        if not t.rows:
            continue
        if header_text in t.cell(0, 0).text:
            return t
        for r in range(len(t.rows)):
            if header_text in t.cell(r, 0).text:
                return t
    return None


def _daily_rows(ws, start, end):
    """Return list of (date, row_index) for dates within [start, end], sorted."""
    out = []
    for row in range(1, ws.max_row + 1):
        d = _as_date(ws.cell(row, 2).value)
        if d and start <= d <= end:
            out.append((d, row))
    out.sort(key=lambda x: x[0])
    return out


# ----------------------------------------------------------------------------
# Sales extraction
# ----------------------------------------------------------------------------
def extract_sales(ws, start, end):
    """Return a dict of 7-day lists for naira, cash, pos, usdt, redot, usd."""
    cols = _find_money_columns(ws)
    rows = _daily_rows(ws, start, end)
    data = {k: [] for k in ["naira", "cash", "pos", "usdt", "redot", "usd"]}
    for _, r in rows:
        data["naira"].append(_num(ws.cell(r, cols.get("NAIRA", 0)).value) if cols.get("NAIRA") else 0.0)
        data["cash"].append(_num(ws.cell(r, cols.get("CASH", 0)).value) if cols.get("CASH") else 0.0)
        data["pos"].append(_num(ws.cell(r, cols.get("POS", 0)).value) if cols.get("POS") else 0.0)
        data["usdt"].append(_num(ws.cell(r, cols.get("USDT", 0)).value) if cols.get("USDT") else 0.0)
        data["redot"].append(_num(ws.cell(r, cols.get("REDOT", 0)).value) if cols.get("REDOT") else 0.0)
        data["usd"].append(_num(ws.cell(r, cols.get("USD", 0)).value) if cols.get("USD") else 0.0)
    # pad to 7 days
    for k in data:
        while len(data[k]) < 7:
            data[k].append(0.0)
    data["_dates"] = [d for d, _ in rows]
    data["_columns_found"] = cols
    return data


# ----------------------------------------------------------------------------
# Expense extraction
# ----------------------------------------------------------------------------
# Map report item label -> spreadsheet item label (when they differ)
EXPENSE_ALIAS = {
    "Goat": "Goat Meat",
    "Mixed Vegetable": "Mix Vegetable",
    "Sweet Pepper Paste": "Sweet pepper paste",
    "Aluminium foil": "Aluminium Foil",
    "Aluminium foil": "Aluminium Foil",
    "Bike Repairs": "Bike Fixing",
    "Delivery Driver": "Kitchen Staff Payment",
    "Electricty": "Electricity",
}


def extract_expenses(ws, start, end, report_items):
    """
    For each report expense item, sum HAMITKOY (Lagos) and MAGUSA (Abuja)
    over the week. Returns {item: (lagos_total, abuja_total)}.
    """
    # locate date columns + their HAMITKOY/MAGUSA sub-columns
    day_cols = []  # (date, hamit_col, magusa_col)
    for col in range(1, ws.max_column + 1):
        d = _as_date(ws.cell(1, col).value)
        if d and start <= d <= end:
            sub_a = str(ws.cell(3, col).value or "").strip().upper()
            sub_b = str(ws.cell(3, col + 1).value or "").strip().upper()
            hamit = col if sub_a == "HAMITKOY" else col
            magusa = col + 1 if sub_b == "MAGUSA" else col + 1
            day_cols.append((d, hamit, magusa))

    # index spreadsheet items by name (column A), rows 6..max
    sheet_items = {}
    for r in range(6, ws.max_row + 1):
        name = ws.cell(r, 1).value
        if name and str(name).strip():
            sheet_items[str(name).strip().lower()] = r

    results = {}
    for item in report_items:
        sheet_name = EXPENSE_ALIAS.get(item, item)
        r = sheet_items.get(sheet_name.strip().lower())
        lagos = abuja = 0.0
        if r:
            for _, hc, mc in day_cols:
                lagos += _num(ws.cell(r, hc).value)
                abuja += _num(ws.cell(r, mc).value)
        results[item] = (lagos, abuja)
    return results


# ----------------------------------------------------------------------------
# Report builder
# ----------------------------------------------------------------------------
def build_report(xlsx_bytes, template_bytes, *, start, end, rates, header, section8=None):
    """
    Fills the UPDATED Weekly Accounting Report template.

    Extraction (HAMITKOY SALES=Lagos, MAGUSA SALES=Abuja, EXPENSE SHEET by item)
    is unchanged. Only the docx-fill half is rewritten to match the new,
    reorganised template:

      T0  header            T1  top summary          T2/T3 Naira day tables
      T4/T5 Cash day tables T6/T7 Day-total tables   T8 crypto amounts
      T9  currency rates    T10 revenue breakdown    T11-14 expense categories
      T15 expense rollup    T16 fixed/operational    T17 net calc   T18 notes

    Fixed/Operational costs (Salaries/Advances, Greep, Other Expenses) and the
    Outstanding Debt figure are manual weekly entries, passed via `section8`.
    """
    wb = openpyxl.load_workbook(io.BytesIO(xlsx_bytes), data_only=True)
    doc = Document(io.BytesIO(template_bytes))
    T = doc.tables
    if len(T) < 19:
        raise ValueError(
            f"This template has {len(T)} tables; the updated Reporta template "
            f"needs 19. Please upload UPDATED_YOOWA_Weekly_Accounting_Report_Template."
        )

    naira_rate = float(rates.get("naira", 32)) or 32.0
    usdt_rate = float(rates.get("usdt", 43))
    usd_rate = float(rates.get("usd", 44.5))
    eur_rate = float(rates.get("eur", 0) or 0)
    gbp_rate = float(rates.get("gbp", 0) or 0)

    lagos = extract_sales(wb["HAMITKOY SALES"], start, end)
    abuja = extract_sales(wb["MAGUSA SALES"], start, end)

    def total(site, key):
        return sum(site[key])

    # ---- T0: header ----
    t = T[0]
    t.cell(0, 1).text = f"{start.strftime('%d/%m/%Y')} to {end.strftime('%d/%m/%Y')}"
    t.cell(0, 3).text = header.get("prepared_by", "FESTUS")
    t.cell(1, 1).text = header.get("exchange_rate_text", f"{int(naira_rate)} = TL1")
    t.cell(1, 3).text = header.get("date_prepared", datetime.date.today().strftime("%d/%m/%Y"))

    # ---- T2/T3: Naira -> TL day tables (Lagos, Abuja) ----
    def fill_naira(table, naira_list):
        tot_n = tot_tl = 0.0
        for i in range(7):
            n = naira_list[i]
            tl = n / naira_rate
            table.cell(i + 1, 1).text = f"{int(round(n)):,}" if n else "0"
            table.cell(i + 1, 2).text = _money(tl)
            tot_n += n
            tot_tl += tl
        table.cell(8, 1).text = f"{int(round(tot_n)):,}"
        table.cell(8, 2).text = _money(tot_tl)
        return tot_tl

    lagos_naira_tl = fill_naira(T[2], lagos["naira"])
    abuja_naira_tl = fill_naira(T[3], abuja["naira"])

    # ---- T4/T5: Cash day tables ----
    def fill_single(table, values):
        tot = 0.0
        for i in range(7):
            table.cell(i + 1, 1).text = _money(values[i])
            tot += values[i]
        table.cell(8, 1).text = _money(tot)
        return tot

    lagos_cash = fill_single(T[4], lagos["cash"])
    abuja_cash = fill_single(T[5], abuja["cash"])

    # ---- Section totals we need for later ----
    lagos_pos = total(lagos, "pos")
    abuja_pos = total(abuja, "pos")
    lagos_usdt = total(lagos, "usdt")
    abuja_usdt = total(abuja, "usdt")
    lagos_redot = total(lagos, "redot")
    abuja_redot = total(abuja, "redot")
    lagos_usd = total(lagos, "usd")
    abuja_usd = total(abuja, "usd")

    # ---- T6/T7: Day-total tables (Naira TL + Cash + POS per day) ----
    def fill_day_totals(table, site):
        tot = 0.0
        for i in range(7):
            day = site["naira"][i] / naira_rate + site["cash"][i] + site["pos"][i]
            table.cell(i + 1, 1).text = _money(day)
            tot += day
        table.cell(8, 1).text = _money(tot)
        return tot

    fill_day_totals(T[6], lagos)
    fill_day_totals(T[7], abuja)

    # ---- T8: crypto amounts (USDT / USD / Redot) in $ ----
    t8 = T[8]
    t8.cell(1, 1).text = f"${lagos_usdt:,.2f}"
    t8.cell(1, 2).text = f"${abuja_usdt:,.2f}  (${lagos_usdt + abuja_usdt:,.2f} combined)"
    t8.cell(2, 1).text = f"${lagos_usd:,.2f}"
    t8.cell(2, 2).text = f"${abuja_usd:,.2f}  (${lagos_usd + abuja_usd:,.2f} combined)"
    t8.cell(3, 1).text = f"${lagos_redot:,.2f}"
    t8.cell(3, 2).text = f"${abuja_redot:,.2f}  (${lagos_redot + abuja_redot:,.2f} combined)"
    sec_l = lagos_usdt + lagos_usd + lagos_redot
    sec_a = abuja_usdt + abuja_usd + abuja_redot
    t8.cell(4, 1).text = f"${sec_l:,.2f}"
    t8.cell(4, 2).text = f"${sec_a:,.2f}  (${sec_l + sec_a:,.2f} combined)"

    # ---- T9: currency rate table (fill Rate + TL Value) ----
    usdt_tl = (lagos_usdt + abuja_usdt) * usdt_rate
    usd_tl = (lagos_usd + abuja_usd) * usd_rate
    t9 = T[9]
    t9.cell(1, 1).text = f"{usdt_rate:g}"
    t9.cell(1, 2).text = _money(usdt_tl)
    t9.cell(2, 1).text = f"{usd_rate:g}"
    t9.cell(2, 2).text = _money(usd_tl)
    if eur_rate:
        t9.cell(3, 1).text = f"{eur_rate:g}"
    if gbp_rate:
        t9.cell(4, 1).text = f"{gbp_rate:g}"

    # TL equivalents used across summary
    usdt_tl_l = lagos_usdt * usdt_rate
    usdt_tl_a = abuja_usdt * usdt_rate
    redot_tl_l = lagos_redot * usdt_rate
    redot_tl_a = abuja_redot * usdt_rate
    foreign_tl_l = lagos_usd * usd_rate
    foreign_tl_a = abuja_usd * usd_rate

    gross_l = (lagos_naira_tl + lagos_cash + lagos_pos + usdt_tl_l
               + redot_tl_l + foreign_tl_l)
    gross_a = (abuja_naira_tl + abuja_cash + abuja_pos + usdt_tl_a
               + redot_tl_a + foreign_tl_a)

    # ---- T10: revenue breakdown ----
    t10 = T[10]

    def rev_row(idx, l, a):
        t10.cell(idx, 1).text = _money(l)
        t10.cell(idx, 2).text = _money(a)
        t10.cell(idx, 3).text = _money(l + a)

    rev_row(1, lagos_naira_tl, abuja_naira_tl)
    rev_row(2, lagos_cash, abuja_cash)
    rev_row(3, lagos_pos, abuja_pos)
    rev_row(4, usdt_tl_l, usdt_tl_a)
    rev_row(5, redot_tl_l, redot_tl_a)
    rev_row(6, foreign_tl_l, foreign_tl_a)
    rev_row(7, gross_l, gross_a)

    # ---- T11-T14: expense category tables (last row = category TOTAL) ----
    def fill_expense_table(table):
        items = [table.cell(r, 0).text.strip() for r in range(1, len(table.rows) - 1)]
        exp = extract_expenses(wb["EXPENSE SHEET"], start, end, items)
        cat_l = cat_a = 0.0
        for i, item in enumerate(items, start=1):
            l, a = exp.get(item, (0.0, 0.0))
            table.cell(i, 1).text = _money(l)
            table.cell(i, 2).text = _money(a)
            table.cell(i, 3).text = _money(l + a)
            cat_l += l
            cat_a += a
        last = len(table.rows) - 1
        table.cell(last, 1).text = _money(cat_l)
        table.cell(last, 2).text = _money(cat_a)
        table.cell(last, 3).text = _money(cat_l + cat_a)
        return cat_l, cat_a

    raw_l, raw_a = fill_expense_table(T[11])      # Raw Materials
    pack_l, pack_a = fill_expense_table(T[12])    # Packaging
    util_l, util_a = fill_expense_table(T[13])    # Utilities
    misc_l, misc_a = fill_expense_table(T[14])    # Miscellaneous

    exp_l = raw_l + pack_l + util_l + misc_l
    exp_a = raw_a + pack_a + util_a + misc_a

    # ---- T15: expense category rollup ----
    t15 = T[15]
    for idx, (l, a) in enumerate(
        [(raw_l, raw_a), (pack_l, pack_a), (util_l, util_a), (misc_l, misc_a)],
        start=1,
    ):
        t15.cell(idx, 1).text = _money(l)
        t15.cell(idx, 2).text = _money(a)
        t15.cell(idx, 3).text = _money(l + a)
    t15.cell(5, 1).text = _money(exp_l)
    t15.cell(5, 2).text = _money(exp_a)
    t15.cell(5, 3).text = _money(exp_l + exp_a)

    # ---- T16: Fixed / Operational costs (manual, from section8) ----
    section8 = section8 or {}
    deductions = section8.get("deductions", [])          # [{label,lagos,abuja}]
    debt_payable = section8.get("debt_payable", "")
    cash_paid = _num(section8.get("cash_payments_made", 0))
    other_misc = _num(section8.get("other_misc", 0))      # e.g. 4700
    other_bulk = _num(section8.get("other_bulk", 0))      # e.g. 15000
    t16 = T[16]
    fix_l = fix_a = 0.0
    body_rows = list(range(1, len(t16.rows) - 2))         # rows above SUBTOTAL
    for ridx, ded in zip(body_rows, deductions):
        label = ded.get("label", "").strip()
        l = _num(ded.get("lagos", 0))
        a = _num(ded.get("abuja", 0))
        if label:
            t16.cell(ridx, 0).text = label
        t16.cell(ridx, 1).text = _money(l)
        t16.cell(ridx, 2).text = _money(a)
        t16.cell(ridx, 3).text = _money(l + a)
        fix_l += l
        fix_a += a
    for ridx in body_rows[len(deductions):]:
        t16.cell(ridx, 1).text = _money(0)
        t16.cell(ridx, 2).text = _money(0)
        t16.cell(ridx, 3).text = _money(0)
    subtotal_row = len(t16.rows) - 2
    t16.cell(subtotal_row, 1).text = _money(fix_l)
    t16.cell(subtotal_row, 2).text = _money(fix_a)
    t16.cell(subtotal_row, 3).text = _money(fix_l + fix_a)
    debt_row = len(t16.rows) - 1
    t16.cell(debt_row, 1).text = "—"
    t16.cell(debt_row, 2).text = "—"
    t16.cell(debt_row, 3).text = f"₺ {debt_payable}" if debt_payable else "—"

    # ---- T17: net calc ----
    net_l = gross_l - exp_l - fix_l
    net_a = gross_a - exp_a - fix_a
    t17 = T[17]

    def net_row(idx, l, a):
        t17.cell(idx, 1).text = _money(l)
        t17.cell(idx, 2).text = _money(a)
        t17.cell(idx, 3).text = _money(l + a)

    net_row(1, gross_l, gross_a)
    net_row(2, exp_l, exp_a)
    net_row(3, fix_l, fix_a)
    net_row(4, net_l, net_a)
    combined_net = net_l + net_a

    # ---- T1: top summary ----
    t1 = T[1]
    net_row_helper = [
        (1, gross_l, gross_a),
        (2, exp_l, exp_a),
        (3, fix_l, fix_a),
        (4, net_l, net_a),
    ]
    for idx, l, a in net_row_helper:
        t1.cell(idx, 1).text = _money(l)
        t1.cell(idx, 2).text = _money(a)
        t1.cell(idx, 3).text = _money(l + a)
    t1.cell(5, 1).text = "—"
    t1.cell(5, 2).text = "—"
    t1.cell(5, 3).text = f"₺ {debt_payable}" if debt_payable else "—"

    # ---- T18: notes ----
    nl = section8.get("notes_lagos", "") or "NIL"
    na = section8.get("notes_abuja", "") or "NIL"
    t18 = T[18]
    t18.cell(0, 0).text = f"Lagos (Lefkosa) Notes: {nl}"
    if len(t18.rows) > 1:
        t18.cell(1, 0).text = f"Abuja (Magusa) Notes: {na}"

    # ---- Auto-filled explanatory notes (wording fixed, numbers live) ----
    other_total = other_misc + other_bulk
    note_updates = {
        "Note: Of the": (
            f"Note: Of the {_money(exp_l)} in Lagos expenses recorded above, "
            f"{_money(cash_paid)} was settled in cash this week; the remaining "
            f"balance forms part of the Outstanding Debt reported in Section 4."
        ),
        "Other Expenses (": (
            f"Other Expenses ({_money(other_total)}, Lagos) includes "
            f"{_money(other_misc)} in general miscellaneous costs and a "
            f"{_money(other_bulk)} bulk chicken stock purchase outside the "
            f"weekly itemized list."
        ),
    }
    for p in doc.paragraphs:
        for prefix, newtext in note_updates.items():
            if p.text.strip().startswith(prefix):
                if p.runs:
                    p.runs[0].text = newtext
                    for r_ in p.runs[1:]:
                        r_.text = ""
                else:
                    p.text = newtext
    
    # ---- Standalone banner paragraphs (overwrite template's sample numbers) ----
    debt_text = f"₺ {debt_payable}" if debt_payable else "—"
    banner_updates = [
        ("COMBINED NET REVENUE BALANCE",
         f"COMBINED NET REVENUE BALANCE = {_money(combined_net)}"),
        ("OUTSTANDING DEBT",
         f"OUTSTANDING DEBT (PAYABLE TO SUPPLIERS) = {debt_text}"),
    ]
    for p in doc.paragraphs:
        for prefix, newtext in banner_updates:
            if p.text.strip().upper().startswith(prefix):
                if p.runs:
                    p.runs[0].text = newtext
                    for r_ in p.runs[1:]:
                        r_.text = ""
                else:
                    p.text = newtext

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read(), {
        "gross_lagos": gross_l, "gross_abuja": gross_a,
        "expenses_lagos": exp_l, "expenses_abuja": exp_a,
        "fixed_lagos": fix_l, "fixed_abuja": fix_a,
        "net_lagos": net_l, "net_abuja": net_a,
        "net": combined_net,
        "days_found": len(lagos["_dates"]),
    }
